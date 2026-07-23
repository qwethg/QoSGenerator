# -*- coding: utf-8 -*-
"""
XML 辅助处理模块
封装所有直接操作 Word 底层 XML 结构的代码，以便与主要的渲染逻辑解耦。
"""
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.enum.text import WD_ALIGN_PARAGRAPH
from services.config_loader import settings

def delete_block_paragraph(paragraph):
    """从文档中物理删除段落的 XML 节点"""
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def delete_block_table(table):
    """从文档中物理删除表格的 XML 节点"""
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_table_by_header(document, header_cells):
    """根据表头内容，在文档中精确查找匹配的表格对象"""
    for table in document.tables:
        if [cell.text.strip() for cell in table.rows[0].cells] == header_cells:
            return table
    return None


def find_paragraph_index(document, expected_text):
    """根据段落的文本内容查找其在文档中的段落索引值"""
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == expected_text:
            return index
    return None


def ensure_table_borders(table):
    """确保表格具备完整的 Word 表格边框（w:tblBorders）"""
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if tblBorders is None:
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr.append(borders_xml)


def format_added_cell(cell, text=''):
    """格式化动态添加的单元格：填充文本、设置居中对齐、垂直居中及显式框线"""
    cell.text = str(text)
    if cell.paragraphs:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vAlign')
    if vAlign is None:
        vAlign_xml = parse_xml(r'<w:vAlign %s w:val="center"/>' % nsdecls('w'))
        tcPr.append(vAlign_xml)

    tcBorders = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
    if tcBorders is None:
        borders_xml = parse_xml(
            r'<w:tcBorders %s>'
            r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(borders_xml)


def apply_imported_table_block(document, config, payload):
    """
    根据前端导入的数据和配置，动态填充或者删除文档中的表格块。
    如果数据启用且存在行，则清空占位行并添加数据；如果禁用，则连同标题、注释物理删除。
    包含表格框线和单元格样式的补齐逻辑。
    """
    title_idx = find_paragraph_index(document, config['title'])
    if title_idx is None:
        return

    target_table = find_table_by_header(document, config['header_cells'])
    if target_table is None:
        return

    title_note = config.get('title_note')
    trailing_note = config.get('trailing_note')

    if not payload['enabled']:
        # 按文本内容查找并删除尾部注释、表格、标题注释、标题
        if trailing_note:
            trailing_idx = find_paragraph_index(document, trailing_note)
            if trailing_idx is not None:
                delete_block_paragraph(document.paragraphs[trailing_idx])
        delete_block_table(target_table)
        if title_note:
            note_idx = find_paragraph_index(document, title_note)
            if note_idx is not None:
                delete_block_paragraph(document.paragraphs[note_idx])
        delete_block_paragraph(document.paragraphs[title_idx])
        return

    # 确保表格主体具有标准边框属性
    ensure_table_borders(target_table)

    # 清空模板数据行（保留第一行表头）
    while len(target_table.rows) > 1:
        target_table._tbl.remove(target_table.rows[1]._tr)

    # 按导入顺序填充数据行，序号自动编号，并应用单元格边框与格式
    for index, row in enumerate(payload['rows'], start=1):
        cells = target_table.add_row().cells
        format_added_cell(cells[0], str(index))
        for cell_index, key in enumerate(config['row_keys'], start=1):
            format_added_cell(cells[cell_index], str(row.get(key, '')).strip())

