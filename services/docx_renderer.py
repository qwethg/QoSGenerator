# -*- coding: utf-8 -*-
"""
Word 文档渲染引擎模块
处理 QoS 审查卡片和互提资料的模板填充、行控制及章节编号重排逻辑。
"""
import os
import io
import re
from docxtpl import DocxTemplate
from docx import Document

from services.config_loader import settings
from services.xml_helper import apply_imported_table_block

# 确定模板根目录（相对于本文件的父目录的父目录）
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TEMPLATE_DIR = os.path.join(BASE_DIR, 'templates_docx')

TEMPLATE_FILE = 'QoS-template.docx'
CROSS_TEMPLATE_FILE = 'htzl_yxtx_sgt.docx'
CROSS_TEMPLATE_FALLBACK_FILES = [
    '互提资料模板-施工图-有线通信提站后 - 副本.docx',
]


def make_rt(text):
    """辅助函数：将多行文本转换为 RichText 以保留换行符"""
    if not text:
        return ''
    from docxtpl import RichText
    rt = RichText()
    # 兼容 \r\n 和 \n
    lines = text.replace('\r\n', '\n').split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            rt.add('\a')  # \a 在 docxtpl 中代表软回车 <w:br/>
        rt.add(line)
    return rt


def safe_bool(val):
    """布尔值安全转换"""
    if isinstance(val, bool):
        return val
    if isinstance(val, str):
        return val.lower() in ('true', '1', 'yes')
    return bool(val)


def format_power(value):
    """用电量格式化：None/空值显示为 —，否则带 kW 单位"""
    if value is None or value == '':
        return '—'
    try:
        return f"{float(value)}kW"
    except (ValueError, TypeError):
        return str(value)


def format_heat(value):
    """散热量格式化：None/空值显示为 —，否则为用电量的 90%"""
    if value is None or value == '':
        return '—'
    try:
        return f"{round(float(value) * 0.9, 1)}kW"
    except (ValueError, TypeError):
        return '—'


def calc_indices(order, enabled_map):
    """根据启用状态自动计算连续序号"""
    idx = 1
    result = {}
    for key in order:
        if enabled_map.get(key, False):
            result[key] = idx
            idx += 1
        else:
            result[key] = ''
    return result


def resolve_cross_template_path():
    """解析互提资料模板路径，兼容本地不同模板文件名。"""
    candidates = [CROSS_TEMPLATE_FILE] + CROSS_TEMPLATE_FALLBACK_FILES
    for filename in candidates:
        path = os.path.join(TEMPLATE_DIR, filename)
        if os.path.exists(path):
            return path
    return None


def build_enabled_chapter_titles(chapter_flags):
    """根据勾选的专业生成连续章节标题映射。"""
    enabled_titles = []
    for flag, title in settings.chapter_title_order:
        if chapter_flags.get(flag, False):
            enabled_titles.append(title)

    title_map = {}
    for idx, title in enumerate(enabled_titles):
        if idx < len(settings.chinese_section_numbers):
            title_map[title] = f"{settings.chinese_section_numbers[idx]}、{title}"
    return title_map


def replace_paragraph_text(paragraph, new_text):
    """替换段落文本并尽量保留首个 run 的样式。"""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new_text)


def renumber_chapter_titles(document, chapter_flags):
    """按启用顺序重排文档中的章节标题序号。"""
    title_map = build_enabled_chapter_titles(chapter_flags)
    if not title_map:
        return

    # 使数字前缀和“、”成为可选，兼容 Word 自动编号将前缀剥离为样式的情形
    title_pattern = re.compile(
        r'^([一二三四五六七八九十]+、\s*)?(?P<title>%s)$'
        % '|'.join(re.escape(title) for _, title in settings.chapter_title_order)
    )

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        match = title_pattern.fullmatch(text)
        if not match:
            continue

        title = match.group('title')
        new_title = title_map.get(title)
        if new_title and text != new_title:
            replace_paragraph_text(paragraph, new_title)


def merge_room_configs(data):
    """合并前端传入的房屋配置与默认配置"""
    rooms = data.get('rooms', {})
    result = {}
    for key, defaults in settings.room_defaults.items():
        user_cfg = rooms.get(key, {})
        cfg = dict(defaults)
        # 允许前端覆盖的字段
        for field in ['enabled', 'name', 'area', 'staff', 'loc', 'power',
                      'power_level', 'voltage', 'power_remark']:
            if field in user_cfg:
                cfg[field] = user_cfg[field]
        # enabled 默认为 True
        cfg['enabled'] = safe_bool(cfg.get('enabled', True))
        result[key] = cfg
    return result


def normalize_imported_table_payload(station_front):
    """规范化导入的表格数据载荷"""
    imported_tables = station_front.get('imported_tables', {}) or {}
    result = {}
    for key in settings.table_block_config:
        payload = imported_tables.get(key, {}) or {}
        rows = payload.get('rows', []) or []
        result[key] = {
            'enabled': safe_bool(payload.get('enabled', False)) and bool(rows),
            'rows': rows if isinstance(rows, list) else [],
        }
    return result


def generate_qos_docx(data):
    """根据 docxtpl 模板和数据生成 QoS docx 文件。"""
    template_path = os.path.join(TEMPLATE_DIR, TEMPLATE_FILE)
    if not os.path.exists(template_path):
        return None, f"模板文件不存在: {TEMPLATE_FILE}"

    try:
        # 1. 加载 docxtpl 模板
        doc = DocxTemplate(template_path)
        
        # 2. 准备上下文数据
        context = {
            'project_name': data.get('project_name', ''),
            'design_range': data.get('design_range', ''),
            'file_name': data.get('file_name', ''),
            'sheet_count': data.get('sheet_count', ''),
            'review_opinions_text': make_rt(data.get('review_opinions_text', '')),
            'audit_suoshi_text': make_rt(data.get('audit_suoshi_text', '')),
            'audit_yuan_text': make_rt(data.get('audit_yuan_text', '')),
            'audit_gongsi_text': make_rt(data.get('audit_gongsi_text', '')),
            'final_audit_level': data.get('final_audit_level', '院级'),
            # 提供默认的复核确认语
            'review_confirmation': data.get('review_confirmation') or '确认已修改。'
        }
        
        # 3. 渲染模板
        doc.render(context)
        
        # 4. 写入内存
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output, None

    except Exception as e:
        import traceback
        return None, traceback.format_exc()


def generate_cross_data_docx(data):
    """根据 docxtpl 模板和数据生成互提资料 docx 文件。"""
    template_path = resolve_cross_template_path()
    if not template_path:
        searched = [CROSS_TEMPLATE_FILE] + CROSS_TEMPLATE_FALLBACK_FILES
        return None, f"互提资料模板不存在: {', '.join(searched)}"

    try:
        doc = DocxTemplate(template_path)

        # 1. 专业章节开关（12 个专业）
        chapter_flags = {
            'has_fangjian': safe_bool(data.get('has_fangjian', False)),
            'has_dianli': safe_bool(data.get('has_dianli', False)),
            'has_nuantong': safe_bool(data.get('has_nuantong', False)),
            'has_qianyin': safe_bool(data.get('has_qianyin', False)),
            'has_jiechuwang': safe_bool(data.get('has_jiechuwang', False)),
            'has_wuxian': safe_bool(data.get('has_wuxian', False)),
            'has_jixie': safe_bool(data.get('has_jixie', False)),
            'has_cheliang': safe_bool(data.get('has_cheliang', False)),
            'has_luji': safe_bool(data.get('has_luji', False)),
            'has_suidao': safe_bool(data.get('has_suidao', False)),
            'has_qiaoliang': safe_bool(data.get('has_qiaoliang', False)),
            'has_zhanchang': safe_bool(data.get('has_zhanchang', False)),
        }

        # 新增通用条款模块逻辑：若选择了新增的4个专业中任意一个，则自动显示通用条款
        chapter_flags['has_common'] = (
            chapter_flags['has_luji'] or
            chapter_flags['has_suidao'] or
            chapter_flags['has_qiaoliang'] or
            chapter_flags['has_zhanchang']
        )

        # 站前基础信息分组映射
        station_front = data.get('station_front', {}) or {}
        
        # 兼容旧版本前端与缓存：如字段缺失则默认开启通信站并用 "XXX" 填充地点
        has_station_txz_val = station_front.get('has_station_txz', None)
        if has_station_txz_val is None:
            has_station_txz = True
        else:
            has_station_txz = safe_bool(has_station_txz_val)
            
        station_txz_loc = station_front.get('station_txz_loc', '')
        if not station_txz_loc:
            station_txz_loc = 'XXX'

        station_front_flags = {
            'has_qlsstdsp': safe_bool(station_front.get('has_qlsstdsp', False)),
            'has_station_txz': has_station_txz,
            'station_txz_loc': station_txz_loc,
        }

        # 电缆槽线路类型方案（多选，归属站前基础信息模块）
        line_schemes = station_front.get('line_schemes', {}) or {}
        line_scheme_flags = {
            'has_line_new_single': safe_bool(line_schemes.get('new_single', False)),
            'has_line_new_double': safe_bool(line_schemes.get('new_double', False)),
            'has_line_add_second': safe_bool(line_schemes.get('add_second', False)),
            'has_line_connecting': safe_bool(line_schemes.get('connecting', False)),
        }

        # 兼容旧版前端字段名
        if 'has_fj' in data:
            chapter_flags['has_fangjian'] = safe_bool(data['has_fj'])
        if 'has_dl' in data:
            chapter_flags['has_dianli'] = safe_bool(data['has_dl'])
        if 'has_cl' in data:
            chapter_flags['has_cheliang'] = safe_bool(data['has_cl'])

        # 2. 合并房屋配置（默认值 + 前端覆盖）
        room_configs = merge_room_configs(data)

        # 3. 电缆沟参数（默认值 + 前端覆盖）
        cable_trench = dict(settings.cable_trench_defaults)
        user_trench = data.get('cable_trench', {})
        for key in cable_trench:
            if key in user_trench:
                try:
                    cable_trench[key] = float(user_trench[key])
                except (ValueError, TypeError):
                    cable_trench[key] = user_trench[key]

        # 4. 计算房屋启用状态
        enabled_map = {key: cfg['enabled'] for key, cfg in room_configs.items()}

        # 5. 自动计算序号
        fj_indices = calc_indices(settings.room_orders['fj'], enabled_map)
        dl_indices = calc_indices(settings.room_orders['dl'], enabled_map)
        me_indices = calc_indices(settings.room_orders['me'], enabled_map)

        # 6. 构建上下文
        context = {
            'project_name': data.get('project_name', ''),
            'design_stage': data.get('design_stage', '施工图'),
            'doc_id': data.get('doc_id', ''),
            'source_institute': data.get('source_institute', '通号院'),
            'source_profession': data.get('source_profession', '有线通信'),
        }
        context.update(chapter_flags)
        context.update(station_front_flags)
        context.update(cable_trench)
        context.update(line_scheme_flags)

        # 7. 映射每个房屋的变量
        for key, cfg in room_configs.items():
            enabled = cfg['enabled']
            context[f'has_{key}'] = enabled
            context[f'area_{key}'] = cfg.get('area', '') if enabled else ''
            context[f'loc_{key}'] = cfg.get('loc', '') if enabled else ''
            context[f'staff_{key}'] = cfg.get('staff', '') if enabled else ''
            context[f'p_{key}'] = format_power(cfg.get('power')) if enabled else '—'
            context[f'h_{key}'] = format_heat(cfg.get('power')) if enabled else '—'
            context[f'idx_{key}'] = fj_indices.get(key, '')
            context[f'idx_el_{key}'] = dl_indices.get(key, '')
            context[f'idx_me_{key}'] = me_indices.get(key, '')

        # 8. 处理其他可能存在的文本字段
        for key, value in data.items():
            if isinstance(value, str) and key not in context:
                context[key] = make_rt(value)

        # 计算导入表格的启用状态，供模板 Jinja2 控制显示/隐藏相关描述
        imported_tables = normalize_imported_table_payload(station_front)
        for block_key, block_val in imported_tables.items():
            context[f'has_{block_key}'] = block_val['enabled']

        doc.render(context)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        rendered_doc = Document(output)
        
        imported_tables = normalize_imported_table_payload(station_front)
        for block_key, block_config in settings.table_block_config.items():
            apply_imported_table_block(rendered_doc, block_config, imported_tables[block_key])

        final_output = io.BytesIO()
        rendered_doc.save(final_output)
        final_output.seek(0)
        return final_output, None

    except Exception as e:
        import traceback
        return None, traceback.format_exc()
