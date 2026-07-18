import docx
import shutil

def replace_text_in_runs(paragraph, old_text, new_text):
    if old_text in paragraph.text:
        full_text = paragraph.text
        new_full_text = full_text.replace(old_text, new_text)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            for run in paragraph.runs:
                run.text = ""
            first_run.text = new_full_text
        else:
            paragraph.add_run(new_full_text)

def modify_template():
    src = 'templates_docx/互提资料模板-施工图.docx.bak'
    dst = 'templates_docx/互提资料模板-施工图.docx'
    
    doc = docx.Document(src)
    
    # 1. Replace project_name, design_stage, doc_id
    for p in doc.paragraphs:
        replace_text_in_runs(p, '新建湛海高铁', '{{ project_name }}')
        replace_text_in_runs(p, '初步设计', '{{ design_stage }}')
        replace_text_in_runs(p, '湛海初通-03', '{{ doc_id }}')

    for p in doc.paragraphs:
        if p.text.startswith('一、房建'):
            p.insert_paragraph_before('{%p if has_fj %}')
        elif p.text.startswith('二、电力'):
            p.insert_paragraph_before('{%p endif %}')
            p.insert_paragraph_before('{%p if has_dl %}')
        elif p.text.startswith('三、暖通'):
            p.insert_paragraph_before('{%p endif %}')
        elif p.text.startswith('八、车辆'):
            p.insert_paragraph_before('{%p if has_cl %}')
    
    doc.add_paragraph('{%p endif %}')

    # 2. Process Tables
    # Table 0: Header
    for row in doc.tables[0].rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                replace_text_in_runs(p, '新建湛海高铁', '{{ project_name }}')
                replace_text_in_runs(p, '初步设计', '{{ design_stage }}')

    t0 = doc.tables[0]
    for row in t0.rows:
        if '房建' in row.cells[1].text and '接收专业' in row.cells[0].text:
            replace_text_in_runs(row.cells[1].paragraphs[0], '房建', '{% if has_fj %}房建{% endif %}')
        elif '电力' in row.cells[1].text and '接收专业' in row.cells[0].text:
            replace_text_in_runs(row.cells[1].paragraphs[0], '电力', '{% if has_dl %}电力{% endif %}')
        elif '车辆' in row.cells[1].text and '接收专业' in row.cells[0].text:
            replace_text_in_runs(row.cells[1].paragraphs[0], '车辆', '{% if has_cl %}车辆{% endif %}')

    # Table 1: fj_rooms
    t1 = doc.tables[1]
    while len(t1.rows) > 2:
        t1._element.remove(t1.rows[-1]._element)
    
    if len(t1.rows) == 2:
        row = t1.rows[1]
        row.cells[0].paragraphs[0].insert_paragraph_before('{%tr for room in fj_rooms %}')
        replace_text_in_runs(row.cells[0].paragraphs[1], row.cells[0].text, '{{ loop.index }}')
        replace_text_in_runs(row.cells[1].paragraphs[0], row.cells[1].text, '{{ room.location }}')
        replace_text_in_runs(row.cells[2].paragraphs[0], row.cells[2].text, '{{ room.name }}')
        replace_text_in_runs(row.cells[3].paragraphs[0], row.cells[3].text, '{{ room.area }}')
        replace_text_in_runs(row.cells[4].paragraphs[0], row.cells[4].text, '{{ room.remark }}')
        row.cells[4].add_paragraph('{%tr endfor %}')

    # Table 3: dl_rooms
    t3 = doc.tables[3]
    while len(t3.rows) > 2:
        t3._element.remove(t3.rows[-1]._element)
    
    if len(t3.rows) == 2:
        row = t3.rows[1]
        row.cells[0].paragraphs[0].insert_paragraph_before('{%tr for room in dl_rooms %}')
        replace_text_in_runs(row.cells[0].paragraphs[1], row.cells[0].text, '{{ loop.index }}')
        replace_text_in_runs(row.cells[1].paragraphs[0], row.cells[1].text, '{{ room.location }}')
        replace_text_in_runs(row.cells[2].paragraphs[0], row.cells[2].text, '{{ room.name }}')
        replace_text_in_runs(row.cells[3].paragraphs[0], row.cells[3].text, '一级') 
        replace_text_in_runs(row.cells[4].paragraphs[0], row.cells[4].text, 'AC 380V')
        replace_text_in_runs(row.cells[5].paragraphs[0], row.cells[5].text, '{{ room.area }}')
        replace_text_in_runs(row.cells[6].paragraphs[0], row.cells[6].text, '{{ room.remark }}')
        row.cells[6].add_paragraph('{%tr endfor %}')

    doc.save(dst)
    print("Template modified successfully!")

if __name__ == '__main__':
    modify_template()