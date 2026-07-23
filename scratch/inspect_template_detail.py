import os
import sys
from docx import Document
from lxml import etree

template_path = r"d:\code\y20_QoSGenerator\templates_docx\htzl_yxtx_sgt.docx"
doc = Document(template_path)

for t_idx in [1, 2, 3]:
    table = doc.tables[t_idx]
    print(f"\n=================== Table {t_idx} ===================")
    print("Row count in template:", len(table.rows))
    for r_idx, row in enumerate(table.rows):
        print(f"--- Row {r_idx} ---")
        for c_idx, cell in enumerate(row.cells):
            tcPr = cell._tc.tcPr
            tcPr_xml = etree.tostring(tcPr, encoding='utf-8').decode('utf-8') if tcPr is not None else "None"
            p_align = cell.paragraphs[0].alignment if cell.paragraphs else None
            print(f" Cell {c_idx} text='{cell.text.strip()}' p_align={p_align} tcPr={tcPr_xml}")
