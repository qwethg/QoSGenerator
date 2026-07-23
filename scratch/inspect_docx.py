import os
from docx import Document
from lxml import etree

template_path = r"d:\code\y20_QoSGenerator\templates_docx\htzl_yxtx_sgt.docx"
doc = Document(template_path)

print(f"Total tables in template: {len(doc.tables)}")

for idx, table in enumerate(doc.tables):
    header = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells] if table.rows else []
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders') if tblPr is not None else None
    tblStyle = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblStyle') if tblPr is not None else None
    style_val = tblStyle.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val') if tblStyle is not None else None
    
    print(f"\n=== Table {idx} ===")
    print(f"Header: {header}")
    print(f"Style: {style_val}")
    print(f"tblBorders present? {tblBorders is not None}")
    if tblBorders is not None:
        print("tblBorders tags:", [child.tag.split('}')[-1] for child in tblBorders])
    if table.rows:
        print(f"Rows count: {len(table.rows)}")
        row_0_tcBorders = [cell._tc.tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders') is not None for cell in table.rows[0].cells if cell._tc.tcPr is not None]
        print(f"Row 0 cells with tcBorders: {row_0_tcBorders}")
