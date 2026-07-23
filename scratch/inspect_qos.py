import os
from docx import Document
from lxml import etree

template_path = r"d:\code\y20_QoSGenerator\templates_docx\QoS-template.docx"
doc = Document(template_path)

print(f"Total tables in QoS-template: {len(doc.tables)}")

for idx, table in enumerate(doc.tables):
    header = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells] if table.rows else []
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders') if tblPr is not None else None
    print(f"\n=== Table {idx} ===")
    print(f"Header: {header}")
    print(f"tblBorders present? {tblBorders is not None}")
