import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from lxml import etree

from services.docx_renderer import generate_cross_data_docx

test_payload = {
    "project_name": "测试工程",
    "has_luji": True,
    "station_front": {
        "imported_tables": {
            "interval_branch_downlead": {
                "enabled": True,
                "rows": [
                    {"mileage": "DK10+100", "remark": "测试备注1"},
                    {"mileage": "DK10+200", "remark": "测试备注2"}
                ]
            },
            "bridge_reserved_downlead": {
                "enabled": True,
                "rows": [
                    {"bridge_name": "特大桥1", "reserved_count": "2", "remark": "备注1"}
                ]
            },
            "cable_crossing_mileage": {
                "enabled": True,
                "rows": [
                    {"mileage": "DK12+300", "remark": "过轨1", "count": "4"}
                ]
            }
        }
    }
}

output, error = generate_cross_data_docx(test_payload)
if error:
    print("Error:", error)
else:
    doc = Document(output)
    print(f"Generated doc with {len(doc.tables)} tables.")
    for idx, table in enumerate(doc.tables):
        header = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells] if table.rows else []
        print(f"\nTable {idx}: header = {header}, rows = {len(table.rows)}")
        tblPr = table._tbl.tblPr
        tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders') if tblPr is not None else None
        print("tblBorders present?", tblBorders is not None)
        if tblBorders is not None:
            print("tblBorders XML:", etree.tostring(tblBorders, encoding='utf-8').decode('utf-8'))
        if len(table.rows) > 1:
            row1_cell0 = table.rows[1].cells[0]
            tcPr = row1_cell0._tc.tcPr
            tcBorders = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders') if tcPr is not None else None
            print("Row 1 Cell 0 tcBorders present?", tcBorders is not None)
