import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from lxml import etree

from services.docx_renderer import generate_cross_data_docx

def ensure_table_borders(table):
    tblPr = table._tbl.tblPr
    tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    if tblBorders is None:
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr.append(borders_xml)

def format_added_cell(cell, text=''):
    cell.text = text
    if cell.paragraphs:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}vAlign')
    if vAlign is None:
        vAlign_xml = parse_xml(r'<w:vAlign %s w:val="center"/>' % nsdecls('w'))
        tcPr.append(vAlign_xml)
        
    tcBorders = tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
    if tcBorders is None:
        borders_xml = parse_xml(
            r'<w:tcBorders %s>'
            r'<w:top w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:left w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:bottom w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'<w:right w:val="single" w:sz="4" w:space="0" w:color="auto"/>'
            r'</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(borders_xml)

full_payload = {
    "project_name": "测试边框修复工程",
    "has_luji": True,
    "station_front": {
        "imported_tables": {
            "interval_branch_downlead": {
                "enabled": True,
                "rows": [
                    {"mileage": "DK10+100", "remark": "测试引下槽1"},
                    {"mileage": "DK10+200", "remark": "测试引下槽2"}
                ]
            },
            "bridge_reserved_downlead": {
                "enabled": True,
                "rows": [
                    {"bridge_name": "特大桥A", "reserved_count": "2", "remark": "备注A"}
                ]
            },
            "cable_crossing_mileage": {
                "enabled": True,
                "rows": [
                    {"mileage": "DK12+300", "remark": "过轨A", "count": "4"}
                ]
            }
        }
    }
}

output, error = generate_cross_data_docx(full_payload)
doc = Document(output)

# Apply fix on doc
for idx in [1, 2, 3]:
    if idx < len(doc.tables):
        tbl = doc.tables[idx]
        ensure_table_borders(tbl)
        for r_idx in range(1, len(tbl.rows)):
            for cell in tbl.rows[r_idx].cells:
                format_added_cell(cell, cell.text)

doc.save(r"d:\code\y20_QoSGenerator\scratch\test_fixed_output.docx")
print("Saved scratch/test_fixed_output.docx")

# Re-inspect doc
doc2 = Document(r"d:\code\y20_QoSGenerator\scratch\test_fixed_output.docx")
for idx in [1, 2, 3]:
    tbl = doc2.tables[idx]
    tblPr = tbl._tbl.tblPr
    tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
    c0 = tbl.rows[1].cells[0]
    tcBorders = c0._tc.tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
    print(f"Table {idx}: tblBorders present? {tblBorders is not None} | row 1 tcBorders present? {tcBorders is not None}")
