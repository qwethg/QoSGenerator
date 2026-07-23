import os
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

from docx import Document
from lxml import etree
from services.docx_renderer import generate_cross_data_docx

full_payload = {
    "project_name": "全专业测试工程",
    "has_fangjian": True,
    "has_dianli": True,
    "has_nuantong": True,
    "has_jixie": True,
    "has_luji": True,
    "has_suidao": True,
    "has_qiaoliang": True,
    "has_zhanchang": True,
    "station_front": {
        "imported_tables": {
            "interval_branch_downlead": {
                "enabled": True,
                "rows": [
                    {"mileage": "DK10+100", "remark": "测试引下槽1"},
                    {"mileage": "DK10+200", "remark": "测试引下槽2"}
                ]
            },
            "bridge_reserved_downlead": {
                "enabled": True,
                "rows": [
                    {"bridge_name": "特大桥A", "reserved_count": "2", "remark": "备注A"}
                ]
            },
            "cable_crossing_mileage": {
                "enabled": True,
                "rows": [
                    {"mileage": "DK12+300", "remark": "过轨A", "count": "4"}
                ]
            }
        }
    }
}

output, error = generate_cross_data_docx(full_payload)
if error:
    print("Error:", error)
else:
    doc = Document(output)
    print(f"Generated doc with {len(doc.tables)} tables.\n")
    for idx, table in enumerate(doc.tables):
        header = [cell.text.strip().replace('\n', ' ') for cell in table.rows[0].cells] if table.rows else []
        tblPr = table._tbl.tblPr
        tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders') if tblPr is not None else None
        
        # Check row 1 (first data row if exists) cell borders
        r1_tcBorders = False
        if len(table.rows) > 1:
            cell0 = table.rows[1].cells[0]
            if cell0._tc.tcPr is not None and cell0._tc.tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders') is not None:
                r1_tcBorders = True

        print(f"Table {idx}: header = {header[:3]}... | total rows = {len(table.rows)}")
        print(f"   tblBorders present: {tblBorders is not None} | row 1 tcBorders present: {r1_tcBorders}")
