# d:\code\y20_QoSGenerator\test_station_front_table_blocks.py
import unittest

from docx import Document

from app import generate_cross_data_docx


def collect_texts(document):
    return [p.text.strip() for p in document.paragraphs if p.text.strip()]


class StationFrontTableBlockTest(unittest.TestCase):
    def test_cable_crossing_block_is_removed_when_not_imported(self):
        data = {
            "project_name": "过轨表块删除测试",
            "has_qiaoliang": True,
            "station_front": {
                "imported_tables": {
                    "cable_crossing_mileage": {
                        "enabled": False,
                        "rows": []
                    }
                }
            }
        }

        output, error = generate_cross_data_docx(data)
        self.assertIsNone(error, msg=error)

        doc = Document(output)
        texts = collect_texts(doc)

        self.assertNotIn("有线通信过轨里程表", texts)
        self.assertNotIn(
            "注：当房屋场坪、隧道救援点等位置变化时，过轨和引下槽里程需配合调整。",
            texts
        )

    def test_cable_crossing_block_is_filled_when_imported(self):
        data = {
            "project_name": "过轨表块填充测试",
            "has_qiaoliang": True,
            "station_front": {
                "imported_tables": {
                    "cable_crossing_mileage": {
                        "enabled": True,
                        "rows": [
                            {"mileage": "DK100+100", "remark": "桥头过轨", "count": "2"},
                            {"mileage": "DK101+800", "remark": "", "count": "4"},
                        ]
                    }
                }
            }
        }

        output, error = generate_cross_data_docx(data)
        self.assertIsNone(error, msg=error)

        doc = Document(output)
        texts = collect_texts(doc)

        self.assertIn("有线通信过轨里程表", texts)
        self.assertIn(
            "注：当房屋场坪、隧道救援点等位置变化时，过轨和引下槽里程需配合调整。",
            texts
        )

        table = next(
            tbl for tbl in doc.tables
            if "根数" in [cell.text.strip() for cell in tbl.rows[0].cells]
        )
        self.assertEqual(table.rows[1].cells[0].text.strip(), "1")
        self.assertEqual(table.rows[1].cells[1].text.strip(), "DK100+100")
        self.assertEqual(table.rows[1].cells[2].text.strip(), "桥头过轨")
        self.assertEqual(table.rows[1].cells[3].text.strip(), "2")
        self.assertEqual(table.rows[2].cells[0].text.strip(), "2")
        self.assertEqual(table.rows[2].cells[1].text.strip(), "DK101+800")
        self.assertEqual(table.rows[2].cells[3].text.strip(), "4")

        # 验证表格边框与单元格边框设置
        tblPr = table._tbl.tblPr
        tblBorders = tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
        self.assertIsNotNone(tblBorders, msg="表格缺少 w:tblBorders 框线定义")

        row1_cell0 = table.rows[1].cells[0]
        tcBorders = row1_cell0._tc.tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
        self.assertIsNotNone(tcBorders, msg="新增单元格缺少 w:tcBorders 显式框线定义")

    def test_interval_branch_downlead_block_is_removed_when_not_imported(self):
        data = {
            "project_name": "分支引下槽删除测试",
            "has_luji": True,
            "station_front": {
                "imported_tables": {
                    "interval_branch_downlead": {
                        "enabled": False,
                        "rows": []
                    }
                }
            }
        }

        output, error = generate_cross_data_docx(data)
        self.assertIsNone(error, msg=error)

        doc = Document(output)
        texts = collect_texts(doc)
        self.assertNotIn("区间分支引下槽里程表", texts)
        self.assertNotIn(
            "（区间无线GSM-R基站、直放站引下槽的里程由无线通信专业计列，具体见无线通信过轨预留里程表）",
            texts
        )

    def test_interval_branch_downlead_block_is_filled_when_imported(self):
        data = {
            "project_name": "区间引下表填充与框线测试",
            "has_luji": True,
            "station_front": {
                "imported_tables": {
                    "interval_branch_downlead": {
                        "enabled": True,
                        "rows": [
                            {"mileage": "DK10+100", "remark": "基站引下槽1"},
                            {"mileage": "DK10+500", "remark": "直放站引下槽2"}
                        ]
                    }
                }
            }
        }

        output, error = generate_cross_data_docx(data)
        self.assertIsNone(error, msg=error)

        doc = Document(output)
        texts = collect_texts(doc)
        self.assertIn("区间分支引下槽里程表", texts)

        table = next(
            tbl for tbl in doc.tables
            if [cell.text.strip() for cell in tbl.rows[0].cells] == ['序号', '里程', '备注']
        )
        self.assertEqual(table.rows[1].cells[0].text.strip(), "1")
        self.assertEqual(table.rows[1].cells[1].text.strip(), "DK10+100")
        self.assertEqual(table.rows[1].cells[2].text.strip(), "基站引下槽1")
        self.assertEqual(table.rows[2].cells[0].text.strip(), "2")
        self.assertEqual(table.rows[2].cells[1].text.strip(), "DK10+500")
        self.assertEqual(table.rows[2].cells[2].text.strip(), "直放站引下槽2")

        # 校验表格框线与单元格框线
        tblBorders = table._tbl.tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
        self.assertIsNotNone(tblBorders, msg="区间引下表缺少 w:tblBorders 框线定义")

        row1_cell0 = table.rows[1].cells[0]
        tcBorders = row1_cell0._tc.tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
        self.assertIsNotNone(tcBorders, msg="区间引下表新增单元格缺少 w:tcBorders 框线")

    def test_bridge_reserved_downlead_block_is_filled_when_imported(self):
        data = {
            "project_name": "桥上预留引下填充测试",
            "has_qiaoliang": True,
            "station_front": {
                "imported_tables": {
                    "bridge_reserved_downlead": {
                        "enabled": True,
                        "rows": [
                            {"bridge_name": "特大桥A", "reserved_count": "2", "remark": "左右线各1处"}
                        ]
                    }
                }
            }
        }

        output, error = generate_cross_data_docx(data)
        self.assertIsNone(error, msg=error)

        doc = Document(output)
        texts = collect_texts(doc)
        self.assertIn("桥上预留引下位置表", texts)

        table = next(
            tbl for tbl in doc.tables
            if [cell.text.strip() for cell in tbl.rows[0].cells] == ['序号', '桥梁', '桥梁引下预留处数', '备注']
        )
        self.assertEqual(table.rows[1].cells[0].text.strip(), "1")
        self.assertEqual(table.rows[1].cells[1].text.strip(), "特大桥A")
        self.assertEqual(table.rows[1].cells[2].text.strip(), "2")
        self.assertEqual(table.rows[1].cells[3].text.strip(), "左右线各1处")

        tblBorders = table._tbl.tblPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tblBorders')
        self.assertIsNotNone(tblBorders, msg="桥上预留引下位置表缺少 w:tblBorders 框线定义")

        row1_cell0 = table.rows[1].cells[0]
        tcBorders = row1_cell0._tc.tcPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcBorders')
        self.assertIsNotNone(tcBorders, msg="桥上预留引下位置表新增单元格缺少 w:tcBorders 框线")

