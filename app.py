# -*- coding: utf-8 -*-
"""
QoS 审查卡片生成器 - Web 应用
基于 Flask 的单页 Web 应用，使用 docxtpl 生成 QoS docx 文件。
"""
import os
import io
import re
from flask import Flask, request, jsonify, send_file, render_template
from docxtpl import DocxTemplate
from docx import Document

app = Flask(__name__, template_folder='templates', static_folder='static')
app.config['TEMPLATES_AUTO_RELOAD'] = True
app.config['SEND_FILE_MAX_AGE_DEFAULT'] = 0
app.jinja_env.auto_reload = True

TEMPLATE_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates_docx')
TEMPLATE_FILE = 'QoS模板.docx'

# 互提资料 - 固定模板
CROSS_TEMPLATE_FILE = '互提资料模板-施工图-有线通信提站后-修复版.docx'
CROSS_TEMPLATE_FALLBACK_FILES = [
    '互提资料模板-施工图-有线通信提站后.docx',
    '互提资料模板-施工图-有线通信提站后 - 副本.docx',
]

CHAPTER_TITLE_ORDER = [
    ('has_common', '通用要求'),
    ('has_luji', '路基专业'),
    ('has_suidao', '隧道专业'),
    ('has_qiaoliang', '桥梁专业'),
    ('has_zhanchang', '站场专业'),
    ('has_fangjian', '房建'),
    ('has_dianli', '电力'),
    ('has_nuantong', '暖通'),
    ('has_qianyin', '牵引变电'),
    ('has_jiechuwang', '接触网'),
    ('has_wuxian', '无线通信'),
    ('has_jixie', '机械'),
    ('has_cheliang', '车辆'),
]

CHINESE_SECTION_NUMBERS = [
    '一', '二', '三', '四', '五', '六', '七', '八', '九', '十',
    '十一', '十二', '十三'
]

# 互提资料 - 房建表房屋顺序
FJ_ROOM_ORDER = [
    'txz', 'xhl_l', 'xhl_m', 'xhl_s', 'xls', 'yrj', 'zf', 'zjz',
    'pcs', 'hyl', 'ds', 'hzgs', 'txcj', 'txgq', 'dlpds', 'dqhst',
    'txyrjx', 'txzbgq', 'txjz'
]

# 互提资料 - 电力表房屋顺序
DL_ROOM_ORDER = [
    'txz', 'xhl_l', 'xhl_m', 'xhl_s', 'xls', 'yrj', 'zf', 'zjz',
    'pcs', 'hyl', 'ds', 'hzgs', 'txcj', 'txgq', 'dlpds', 'dqhst',
    'txyrjx', 'txiz'
]

# 互提资料 - 机械表房屋顺序
ME_ROOM_ORDER = ['txcj', 'txgq']

# 互提资料 - 房屋默认参数（面积、定员、用电量、供电等级、电压、备注）
ROOM_DEFAULTS = {
    'txz': {
        'name': '通信站',
        'area': '1900', 'staff': '15', 'loc': '枢纽所在地',
        'power': 100, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '',
    },
    'xhl_l': {
        'name': '信号楼（大）',
        'area': '120', 'staff': '', 'loc': '车站',
        'power': 45, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '枢纽重要站点',
    },
    'xhl_m': {
        'name': '信号楼（中）',
        'area': '100', 'staff': '', 'loc': '车站',
        'power': 25, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '视频节点、市级',
    },
    'xhl_s': {
        'name': '信号楼（小）',
        'area': '80', 'staff': '', 'loc': '车站',
        'power': 15, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '其他',
    },
    'xls': {
        'name': '线路所通信机械室',
        'area': '80', 'staff': '', 'loc': '线路所',
        'power': 25, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '',
    },
    'yrj': {
        'name': '通信引入间',
        'area': '15', 'staff': '', 'loc': '引入间',
        'power': 5, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '',
    },
    'zf': {
        'name': '站房通信机械室',
        'area': '80', 'staff': '', 'loc': '站房',
        'power': 15, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '不含信号楼的站房',
    },
    'zjz': {
        'name': '信号中继站通信机械室',
        'area': '60', 'staff': '', 'loc': '中继站',
        'power': 15, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '',
    },
    'pcs': {
        'name': '公安派出所通信机械室',
        'area': '30', 'staff': '', 'loc': '公安派出所',
        'power': 10, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '',
    },
    'hyl': {
        'name': '货运楼通信机械室',
        'area': '60', 'staff': '', 'loc': '货运楼',
        'power': 15, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '',
    },
    'ds': {
        'name': '动车所、综合维修段等通信机械室',
        'area': '80', 'staff': '', 'loc': '动车所/综合维修段',
        'power': 15, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '',
    },
    'hzgs': {
        'name': '单身宿舍、办公楼配线间',
        'area': '10', 'staff': '', 'loc': '单身宿舍/办公楼',
        'power': 2, 'power_level': '—', 'voltage': 'AC 220V±10%',
        'power_remark': '公司管理',
    },
    'txcj': {
        'name': '存车场生产办公楼通信机械室',
        'area': '60', 'staff': '', 'loc': '存车场',
        'power': 15, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '',
    },
    'txgq': {
        'name': '综合维修工区办公楼通信机械室',
        'area': '60', 'staff': '', 'loc': '综合维修工区',
        'power': 15, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '',
    },
    'dlpds': {
        'name': '电力配电所通信机械室',
        'area': '30', 'staff': '', 'loc': '电力配电所',
        'power': 10, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '箱式机房设于主控室',
    },
    'dqhst': {
        'name': '电气化所亭通信机械室',
        'area': '30', 'staff': '', 'loc': '电气化所亭',
        'power': 10, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '箱式机房设于主控室',
    },
    'txyrjx': {
        'name': '车辆探测站（5T机房）',
        'area': '10', 'staff': '', 'loc': '车辆探测站',
        'power': 2, 'power_level': '—', 'voltage': 'AC 220V±10%',
        'power_remark': '通信设备用插座',
    },
    'txzbgq': {
        'name': '未设工区的地点通信值班工区',
        'area': '10', 'staff': '', 'loc': '通信值班工区',
        'power': 2, 'power_level': '—', 'voltage': 'AC 220V±10%',
        'power_remark': '',
    },
    'txiz': {
        'name': '区间基站通信机械室',
        'area': '', 'staff': '', 'loc': '区间基站',
        'power': None, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '由无线专业一并提出',
    },
    'txjz': {
        'name': '区间基站通信机械室（房建）',
        'area': '', 'staff': '', 'loc': '区间基站',
        'power': None, 'power_level': '一级', 'voltage': 'AC 380V±10%',
        'power_remark': '由无线专业一并提出',
    },
}

# 电缆沟/分支槽默认值（单位：mm，间距：m）
CABLE_TRENCH_DEFAULTS = {
    'trench_txz_width': 500,
    'trench_txz_depth': 400,
    'trench_mid_width': 400,
    'trench_mid_depth': 400,
    'branch_dist_min': 2,
    'branch_txz_width': 500,
    'branch_txz_depth': 400,
    'branch_major_width': 400,
    'branch_major_depth': 400,
    'trough_width': 250,
    'trough_depth': 150,
}


def make_rt(text):
    """辅助函数：将多行文本转换为 RichText 以保留换行符"""
    if not text:
        return ''
    from docxtpl import RichText
    rt = RichText()
    # 兼容 \r\n 和 \n
    lines = text.replace('\r\n', '\n').split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            rt.add('\a')  # \a 在 docxtpl 中代表软回车 <w:br/>
        rt.add(line)
    return rt


def safe_bool(val):
    """布尔值安全转换"""
    if isinstance(val, bool):
        return val
    if isinstance(val, str):
        return val.lower() in ('true', '1', 'yes')
    return bool(val)


def format_power(value):
    """用电量格式化：None/空值显示为 —，否则带 kW 单位"""
    if value is None or value == '':
        return '—'
    try:
        return f"{float(value)}kW"
    except (ValueError, TypeError):
        return str(value)


def format_heat(value):
    """散热量格式化：None/空值显示为 —，否则为用电量的 90%"""
    if value is None or value == '':
        return '—'
    try:
        return f"{round(float(value) * 0.9, 1)}kW"
    except (ValueError, TypeError):
        return '—'


def calc_indices(order, enabled_map):
    """根据启用状态自动计算连续序号"""
    idx = 1
    result = {}
    for key in order:
        if enabled_map.get(key, False):
            result[key] = idx
            idx += 1
        else:
            result[key] = ''
    return result


def resolve_cross_template_path():
    """解析互提资料模板路径，兼容本地不同模板文件名。"""
    candidates = [CROSS_TEMPLATE_FILE] + CROSS_TEMPLATE_FALLBACK_FILES
    for filename in candidates:
        path = os.path.join(TEMPLATE_DIR, filename)
        if os.path.exists(path):
            return path
    return None


def build_enabled_chapter_titles(chapter_flags):
    """根据勾选的专业生成连续章节标题映射。"""
    enabled_titles = []
    for flag, title in CHAPTER_TITLE_ORDER:
        if chapter_flags.get(flag, False):
            enabled_titles.append(title)

    title_map = {}
    for idx, title in enumerate(enabled_titles):
        if idx < len(CHINESE_SECTION_NUMBERS):
            title_map[title] = f"{CHINESE_SECTION_NUMBERS[idx]}、{title}"
    return title_map


def replace_paragraph_text(paragraph, new_text):
    """替换段落文本并尽量保留首个 run 的样式。"""
    if paragraph.runs:
        paragraph.runs[0].text = new_text
        for run in paragraph.runs[1:]:
            run.text = ''
    else:
        paragraph.add_run(new_text)


def renumber_chapter_titles(document, chapter_flags):
    """按启用顺序重排文档中的章节标题序号。"""
    title_map = build_enabled_chapter_titles(chapter_flags)
    if not title_map:
        return

    title_pattern = re.compile(
        r'^(?P<num>[一二三四五六七八九十]+)、\s*(?P<title>%s)$'
        % '|'.join(re.escape(title) for _, title in CHAPTER_TITLE_ORDER)
    )

    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        match = title_pattern.fullmatch(text)
        if not match:
            continue

        title = match.group('title')
        new_title = title_map.get(title)
        if new_title and text != new_title:
            replace_paragraph_text(paragraph, new_title)


def merge_room_configs(data):
    """合并前端传入的房屋配置与默认配置"""
    rooms = data.get('rooms', {})
    result = {}
    for key, defaults in ROOM_DEFAULTS.items():
        user_cfg = rooms.get(key, {})
        cfg = dict(defaults)
        # 允许前端覆盖的字段
        for field in ['enabled', 'name', 'area', 'staff', 'loc', 'power',
                      'power_level', 'voltage', 'power_remark']:
            if field in user_cfg:
                cfg[field] = user_cfg[field]
        # enabled 默认为 True
        cfg['enabled'] = safe_bool(cfg.get('enabled', True))
        result[key] = cfg
    return result


def generate_qos_docx(data):
    """根据 docxtpl 模板和数据生成 QoS docx 文件。"""
    template_path = os.path.join(TEMPLATE_DIR, TEMPLATE_FILE)
    if not os.path.exists(template_path):
        return None, f'模板文件不存在: {TEMPLATE_FILE}'

    try:
        from docxtpl import DocxTemplate
        
        # 1. 加载 docxtpl 模板
        doc = DocxTemplate(template_path)
        
        # 2. 准备上下文数据
        context = {
            'project_name': data.get('project_name', ''),
            'design_range': data.get('design_range', ''),
            'file_name': data.get('file_name', ''),
            'sheet_count': data.get('sheet_count', ''),
            'review_opinions_text': make_rt(data.get('review_opinions_text', '')),
            'audit_suoshi_text': make_rt(data.get('audit_suoshi_text', '')),
            'audit_yuan_text': make_rt(data.get('audit_yuan_text', '')),
            'audit_gongsi_text': make_rt(data.get('audit_gongsi_text', '')),
            'final_audit_level': data.get('final_audit_level', '院级'),
            # 提供默认的复核确认语
            'review_confirmation': data.get('review_confirmation') or '确认已修改。'
        }
        
        # 3. 渲染模板
        doc.render(context)
        
        # 4. 写入内存
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        return output, None

    except Exception as e:
        import traceback
        return None, traceback.format_exc()


TABLE_BLOCK_CONFIG = {
    'interval_branch_downlead': {
        'title': '区间分支引下槽里程表',
        'header_cells': ['序号', '里程', '备注'],
        'title_note': '（区间无线GSM-R基站、直放站引下槽的里程由无线通信专业计列，具体见无线通信过轨预留里程表）',
        'trailing_note': None,
        'row_keys': ['mileage', 'remark'],
    },
    'bridge_reserved_downlead': {
        'title': '桥上预留引下位置表',
        'header_cells': ['序号', '桥梁', '桥梁引下预留处数', '备注'],
        'title_note': None,
        'trailing_note': None,
        'row_keys': ['bridge_name', 'reserved_count', 'remark'],
    },
    'cable_crossing_mileage': {
        'title': '有线通信过轨里程表',
        'header_cells': ['序号', '里程', '备   注', '根数'],
        'title_note': '（区间无线GSM-R基站、直放站过轨已由无线通信专业计列，具体见无线区间设施设置里程表；不包含双线隧道综合洞室过轨里程，双线隧道综合洞室里程见隧道专业综合洞室里程表）',
        'trailing_note': '注：当房屋场坪、隧道救援点等位置变化时，过轨和引下槽里程需配合调整。',
        'row_keys': ['mileage', 'remark', 'count'],
    },
}


def normalize_imported_table_payload(station_front):
    imported_tables = station_front.get('imported_tables', {}) or {}
    result = {}
    for key in TABLE_BLOCK_CONFIG:
        payload = imported_tables.get(key, {}) or {}
        rows = payload.get('rows', []) or []
        result[key] = {
            'enabled': safe_bool(payload.get('enabled', False)) and bool(rows),
            'rows': rows if isinstance(rows, list) else [],
        }
    return result


def delete_block_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def delete_block_table(table):
    element = table._element
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)


def find_table_by_header(document, header_cells):
    for table in document.tables:
        if [cell.text.strip() for cell in table.rows[0].cells] == header_cells:
            return table
    return None


def find_paragraph_index(document, expected_text):
    for index, paragraph in enumerate(document.paragraphs):
        if paragraph.text.strip() == expected_text:
            return index
    return None


def apply_imported_table_block(document, config, payload):
    title_idx = find_paragraph_index(document, config['title'])
    if title_idx is None:
        return

    target_table = find_table_by_header(document, config['header_cells'])
    if target_table is None:
        return

    title_note = config.get('title_note')
    trailing_note = config.get('trailing_note')

    if not payload['enabled']:
        # 按文本内容查找并删除尾部注释、表格、标题注释、标题
        if trailing_note:
            trailing_idx = find_paragraph_index(document, trailing_note)
            if trailing_idx is not None:
                delete_block_paragraph(document.paragraphs[trailing_idx])
        delete_block_table(target_table)
        if title_note:
            note_idx = find_paragraph_index(document, title_note)
            if note_idx is not None:
                delete_block_paragraph(document.paragraphs[note_idx])
        delete_block_paragraph(document.paragraphs[title_idx])
        return

    # 清空模板数据行（保留表头）
    while len(target_table.rows) > 1:
        target_table._tbl.remove(target_table.rows[1]._tr)

    # 按导入顺序填充数据行，序号自动编号
    for index, row in enumerate(payload['rows'], start=1):
        cells = target_table.add_row().cells
        cells[0].text = str(index)
        for cell_index, key in enumerate(config['row_keys'], start=1):
            cells[cell_index].text = str(row.get(key, '')).strip()


def generate_cross_data_docx(data):
    """根据 docxtpl 模板和数据生成互提资料 docx 文件。"""
    template_path = resolve_cross_template_path()
    if not template_path:
        searched = [CROSS_TEMPLATE_FILE] + CROSS_TEMPLATE_FALLBACK_FILES
        return None, f"互提资料模板不存在: {', '.join(searched)}"

    try:
        doc = DocxTemplate(template_path)

        # 1. 专业章节开关（12 个专业）
        chapter_flags = {
            'has_fangjian': safe_bool(data.get('has_fangjian', False)),
            'has_dianli': safe_bool(data.get('has_dianli', False)),
            'has_nuantong': safe_bool(data.get('has_nuantong', False)),
            'has_qianyin': safe_bool(data.get('has_qianyin', False)),
            'has_jiechuwang': safe_bool(data.get('has_jiechuwang', False)),
            'has_wuxian': safe_bool(data.get('has_wuxian', False)),
            'has_jixie': safe_bool(data.get('has_jixie', False)),
            'has_cheliang': safe_bool(data.get('has_cheliang', False)),
            'has_luji': safe_bool(data.get('has_luji', False)),
            'has_suidao': safe_bool(data.get('has_suidao', False)),
            'has_qiaoliang': safe_bool(data.get('has_qiaoliang', False)),
            'has_zhanchang': safe_bool(data.get('has_zhanchang', False)),
        }

        # 新增通用条款模块逻辑：若选择了新增的4个专业中任意一个，则自动显示通用条款
        chapter_flags['has_common'] = (
            chapter_flags['has_luji'] or
            chapter_flags['has_suidao'] or
            chapter_flags['has_qiaoliang'] or
            chapter_flags['has_zhanchang']
        )

        # 站前基础信息分组映射
        station_front = data.get('station_front', {}) or {}
        station_front_flags = {
            'has_qlsstdsp': safe_bool(station_front.get('has_qlsstdsp', False)),
        }

        # 兼容旧版前端字段名
        if 'has_fj' in data:
            chapter_flags['has_fangjian'] = safe_bool(data['has_fj'])
        if 'has_dl' in data:
            chapter_flags['has_dianli'] = safe_bool(data['has_dl'])
        if 'has_cl' in data:
            chapter_flags['has_cheliang'] = safe_bool(data['has_cl'])

        # 2. 合并房屋配置（默认值 + 前端覆盖）
        room_configs = merge_room_configs(data)

        # 3. 电缆沟参数（默认值 + 前端覆盖）
        cable_trench = dict(CABLE_TRENCH_DEFAULTS)
        user_trench = data.get('cable_trench', {})
        for key in cable_trench:
            if key in user_trench:
                try:
                    cable_trench[key] = float(user_trench[key])
                except (ValueError, TypeError):
                    cable_trench[key] = user_trench[key]

        # 4. 计算房屋启用状态
        enabled_map = {key: cfg['enabled'] for key, cfg in room_configs.items()}

        # 5. 自动计算序号
        fj_indices = calc_indices(FJ_ROOM_ORDER, enabled_map)
        dl_indices = calc_indices(DL_ROOM_ORDER, enabled_map)
        me_indices = calc_indices(ME_ROOM_ORDER, enabled_map)

        # 6. 构建上下文
        context = {
            'project_name': data.get('project_name', ''),
            'design_stage': data.get('design_stage', '施工图'),
            'doc_id': data.get('doc_id', ''),
            'source_institute': data.get('source_institute', '通号院'),
            'source_profession': data.get('source_profession', '有线通信'),
        }
        context.update(chapter_flags)
        context.update(station_front_flags)
        context.update(cable_trench)

        # 7. 映射每个房屋的变量
        for key, cfg in room_configs.items():
            enabled = cfg['enabled']
            context[f'has_{key}'] = enabled
            context[f'area_{key}'] = cfg.get('area', '') if enabled else ''
            context[f'loc_{key}'] = cfg.get('loc', '') if enabled else ''
            context[f'staff_{key}'] = cfg.get('staff', '') if enabled else ''
            context[f'p_{key}'] = format_power(cfg.get('power')) if enabled else '—'
            context[f'h_{key}'] = format_heat(cfg.get('power')) if enabled else '—'
            context[f'idx_{key}'] = fj_indices.get(key, '')
            context[f'idx_el_{key}'] = dl_indices.get(key, '')
            context[f'idx_me_{key}'] = me_indices.get(key, '')

        # 8. 处理其他可能存在的文本字段
        for key, value in data.items():
            if isinstance(value, str) and key not in context:
                context[key] = make_rt(value)

        doc.render(context)

        output = io.BytesIO()
        doc.save(output)
        output.seek(0)

        rendered_doc = Document(output)
        renumber_chapter_titles(rendered_doc, chapter_flags)

        imported_tables = normalize_imported_table_payload(station_front)
        for block_key, config in TABLE_BLOCK_CONFIG.items():
            apply_imported_table_block(rendered_doc, config, imported_tables[block_key])

        final_output = io.BytesIO()
        rendered_doc.save(final_output)
        final_output.seek(0)
        return final_output, None

    except Exception as e:
        import traceback
        return None, traceback.format_exc()


@app.route('/', defaults={'path': ''})
@app.route('/<path:path>')
def catch_all(path):
    # 如果请求的是 API，就不走这里
    if path.startswith('api/'):
        return "Not Found", 404
    return render_template('index.html')


@app.route('/api/generate/<doc_type>', methods=['POST'])
def api_generate(doc_type):
    data = request.json
    
    if doc_type == 'qos':
        output, error = generate_qos_docx(data)
        filename = f"QoS-{data.get('project_name', 'output')}.docx"
    elif doc_type == 'cross_data':
        output, error = generate_cross_data_docx(data)
        filename = f"互提资料-{data.get('project_name', 'output')}.docx"
    else:
        return jsonify({'error': '未知的文档类型'}), 400

    if error:
        return jsonify({'error': error}), 500

    return send_file(
        output,
        as_attachment=True,
        download_name=filename,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


if __name__ == '__main__':
    # Vercel/容器部署：通过 PORT 环境变量注入端口；本地默认 5000
    port = int(os.environ.get('PORT', 5000))
    os.makedirs(TEMPLATE_DIR, exist_ok=True)
    print("=" * 50)
    print("QoS 审查卡片生成器 (docxtpl 实验分支)")
    print("=" * 50)
    print(f"模板文件: {TEMPLATE_FILE}")
    print(f"模板目录: {TEMPLATE_DIR}")
    print(f"访问地址: http://localhost:{port}")
    print("=" * 50)
    app.run(host='0.0.0.0', port=port, debug=False)
