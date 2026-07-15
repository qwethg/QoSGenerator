# -*- coding: utf-8 -*-
"""
自动为模板中每个 {%tr if has_xxx %} 行补充 {%tr endif %}
跳过垂直合并产生的重复行
"""
import shutil
import docx

src = 'templates_docx/互提资料模板-施工图-有线通信提站后.docx'
bak = 'templates_docx/互提资料模板-施工图-有线通信提站后.docx.bak'

# 备份
shutil.copy2(src, bak)
print('已备份到', bak)

doc = docx.Document(src)
fixed_count = 0

for ti, table in enumerate(doc.tables):
    prev_tc_refs = None
    for ri, row in enumerate(table.rows):
        row_text = ''.join(cell.text for cell in row.cells)

        # 只处理包含 {%tr if 且没有 {%tr endif 的行
        if '{%tr' not in row_text or 'if' not in row_text or 'endif' in row_text:
            prev_tc_refs = [cell._tc for cell in row.cells]
            continue

        # 检查是否是垂直合并的重复行（_tc 元素与上一行完全相同）
        current_tc_refs = [cell._tc for cell in row.cells]
        if prev_tc_refs is not None and len(current_tc_refs) == len(prev_tc_refs):
            is_duplicate = all(a is b for a, b in zip(current_tc_refs, prev_tc_refs))
            if is_duplicate:
                prev_tc_refs = current_tc_refs
                continue

        # 在最后一个单元格的最后一个段落末尾添加 {%tr endif %}
        last_cell = row.cells[-1]
        if last_cell.paragraphs:
            last_para = last_cell.paragraphs[-1]
        else:
            last_para = last_cell.add_paragraph()
        last_para.add_run(' {%tr endif %}')
        fixed_count += 1
        print('  表格%d 行%d: 已补充 endif' % (ti, ri))

        prev_tc_refs = current_tc_refs

doc.save(src)
print('\n共补充 %d 个 {%tr endif %}' % fixed_count)
